"""
Creates HOW_TO_USE.docx — a formatted Word document with step-by-step
instructions for using the NVH Pre/Post BCW Report Generator.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = os.path.join(os.path.dirname(os.path.dirname(__file__)),
                      "HOW_TO_USE.docx")

# ── Colour palette ────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x4E, 0x79)
MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF0)
GOLD       = RGBColor(0xBF, 0x8F, 0x00)
GOLD_LIGHT = RGBColor(0xFF, 0xF2, 0xCC)
GREEN_DARK = RGBColor(0x37, 0x5C, 0x23)
GREEN_LIGHT= RGBColor(0xE2, 0xEF, 0xDA)
RED        = RGBColor(0xC0, 0x00, 0x00)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY_TEXT  = RGBColor(0x59, 0x59, 0x59)
BLACK      = RGBColor(0x00, 0x00, 0x00)


def set_cell_bg(cell, rgb: RGBColor):
    """Set background colour of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_colour = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def set_cell_borders(cell):
    """Add thin borders to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "BFBFBF")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1, colour=DARK_BLUE, space_before=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    run.font.color.rgb = colour
    return p


def add_body(doc, text, bold=False, italic=False, colour=BLACK,
             space_before=2, space_after=2, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.bold             = bold
    run.italic           = italic
    run.font.size        = Pt(11)
    run.font.color.rgb   = colour
    return p


def add_step_header(doc, number, title, colour=DARK_BLUE):
    """Adds a visually distinct step banner."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run_num = p.add_run(f"STEP {number}  ")
    run_num.bold           = True
    run_num.font.size      = Pt(13)
    run_num.font.color.rgb = colour
    run_title = p.add_run(f"— {title}")
    run_title.bold           = True
    run_title.font.size      = Pt(13)
    run_title.font.color.rgb = colour
    # Bottom border under the step header
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), f"{colour[0]:02X}{colour[1]:02X}{colour[2]:02X}")
    pBdr.append(bot)
    pPr.append(pBdr)


def add_numbered_step(doc, number, text, bold_part=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.6)
    if bold_part and bold_part in text:
        before, after = text.split(bold_part, 1)
        if before:
            r = p.add_run(before); r.font.size = Pt(11)
        rb = p.add_run(bold_part); rb.bold = True; rb.font.size = Pt(11)
        if after:
            r = p.add_run(after); r.font.size = Pt(11)
    else:
        r = p.add_run(text); r.font.size = Pt(11)


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.8)
    run = p.add_run(f"Tip:  {text}")
    run.italic         = True
    run.font.size      = Pt(10)
    run.font.color.rgb = GREY_TEXT


def add_two_col_table(doc, rows, header_bg=DARK_BLUE, row_bg=LIGHT_BLUE):
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Cm(7)
    table.columns[1].width = Cm(9)

    # Header row
    for j, h in enumerate(["File / Item", "What It Is"]):
        cell = table.cell(0, j)
        set_cell_bg(cell, header_bg)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold           = True
        run.font.color.rgb = WHITE
        run.font.size      = Pt(11)

    # Data rows
    for i, (col1, col2) in enumerate(rows, start=1):
        for j, text in enumerate([col1, col2]):
            cell = table.cell(i, j)
            set_cell_bg(cell, LIGHT_BLUE if i % 2 == 0 else WHITE)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10)
            if j == 0:
                run.bold = True


def add_pf_table(doc):
    rows = [
        ("PASS",             "dB value is +3 or below  (all negative values are PASS)",
         GREEN_DARK, GREEN_LIGHT),
        ("CONDITIONAL FAIL", "dB value is between +3 and +6",
         GOLD,       GOLD_LIGHT),
        ("ABSOLUTE FAIL",    "dB value is above +6",
         RED,        RGBColor(0xFF, 0xE0, 0xE0)),
    ]
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)

    for j, h in enumerate(["Result", "When It Applies"]):
        cell = table.cell(0, j)
        set_cell_bg(cell, DARK_BLUE)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(11)

    for i, (result, condition, text_col, bg_col) in enumerate(rows, start=1):
        for j, text in enumerate([result, condition]):
            cell = table.cell(i, j)
            set_cell_bg(cell, bg_col)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold           = (j == 0)
            run.font.color.rgb = text_col if j == 0 else BLACK
            run.font.size      = Pt(11)


# ── Build the document ────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Title ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("NVH Pre/Post BCW Report Generator")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = DARK_BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(16)
r2 = p2.add_run("Step-by-Step Instructions")
r2.font.size = Pt(13); r2.font.color.rgb = GREY_TEXT; r2.italic = True

doc.add_paragraph()  # spacer

# ── What this tool does ───────────────────────────────────────────────
add_heading(doc, "What This Tool Does", level=2, space_before=0)
add_body(doc,
    "This tool automates the repetitive parts of building NVH Pre/Post BCW "
    "durability reports in PowerPoint. For each new test campaign it:")
for bullet in [
    "Generates one PowerPoint report per sample with all test information filled in automatically",
    "After your engineering review, fills in the dB values and Pass/Fail results automatically",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(bullet); r.font.size = Pt(11)

add_body(doc,
    "You only need to do the engineering work — the repetitive data entry is handled for you.",
    bold=True, colour=MID_BLUE, space_before=6)

doc.add_paragraph()

# ── Files you will use ────────────────────────────────────────────────
add_heading(doc, "Files You Will Use", level=2)
add_two_col_table(doc, [
    ("working/report_template.pptx",        "Your blank template — prepared once per campaign"),
    ("working/NVH_Report_Input.xlsx",        "Your input form — fill in test and sample data here"),
    ("working/Step2_Generate_Reports.bat",   "Double-click to generate all reports"),
    ("working/Step7_Update_Results.bat",     "Double-click to update reports with results"),
    ("reports/",                             "Folder where your completed reports appear"),
])

doc.add_paragraph()

# ── Step 1 ─────────────────────────────────────────────────────────────
add_step_header(doc, 1, "Prepare Your Template", colour=DARK_BLUE)
add_body(doc, "Do this once at the start of each campaign.", italic=True, colour=GREY_TEXT)
add_numbered_step(doc, 1, "Open  working/report_template.pptx  in PowerPoint",
                  bold_part="working/report_template.pptx")
add_numbered_step(doc, 2, "Find every dB result value on every slide  (e.g. -23dB, -30dB, -16dB)")
add_numbered_step(doc, 3, "Replace each one by typing exactly  XXdB  — no spaces, capital XX",
                  bold_part="XXdB")
add_numbered_step(doc, 4, "Save the file — keep the same filename and location")
add_numbered_step(doc, 5, "Close PowerPoint")
add_note(doc, "The template already contains a copy of your previous report as a starting point.")

# ── Step 2 ─────────────────────────────────────────────────────────────
add_step_header(doc, 2, "Fill In Your Test Information", colour=DARK_BLUE)
add_body(doc, "Do this before generating reports.", italic=True, colour=GREY_TEXT)
add_numbered_step(doc, 1, "Open  working/NVH_Report_Input.xlsx",
                  bold_part="working/NVH_Report_Input.xlsx")
add_numbered_step(doc, 2,
    "Go to the Campaign_Info sheet — fill in the programme and test details  "
    "(programme name, test order numbers, test dyno, dates, part ratio, etc.)",
    bold_part="Campaign_Info")
add_numbered_step(doc, 3,
    "Go to the Samples sheet — fill in one row per sample (blue columns only)  "
    "(part number, serial number, sample number, published date)",
    bold_part="Samples")
add_numbered_step(doc, 4, "Save and close the Excel file")
add_numbered_step(doc, 5, "Double-click  Step2_Generate_Reports.bat",
                  bold_part="Step2_Generate_Reports.bat")
add_numbered_step(doc, 6, "A window will appear showing progress — wait for it to finish")
add_numbered_step(doc, 7, "Your reports appear in the  reports/  folder — one per sample",
                  bold_part="reports/")

# ── Steps 3-6 ──────────────────────────────────────────────────────────
add_step_header(doc, "3 to 6", "Your Engineering Work  (manual)", colour=MID_BLUE)
add_body(doc, "This part remains manual — it requires your engineering expertise.",
         italic=True, colour=GREY_TEXT)
for bullet in [
    "Open each report from the  reports/  folder",
    "Replace the ActivePictures with the correct TestLab plots",
    "Format the ActivePictures to match your standard layout",
    "Review each plot carefully",
    "Fill in the actual dB values directly in the PowerPoint slides",
    "Add any comments",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(bullet); r.font.size = Pt(11)

# ── Step 7 ─────────────────────────────────────────────────────────────
add_step_header(doc, 7, "Fill In Results and Update Pass/Fail", colour=DARK_BLUE)
add_body(doc, "Do this after completing your engineering review.", italic=True, colour=GREY_TEXT)
add_numbered_step(doc, 1, "Open  working/NVH_Report_Input.xlsx",
                  bold_part="working/NVH_Report_Input.xlsx")
add_numbered_step(doc, 2,
    "Go to the Samples sheet — fill in the dB values in the green (Step 7) columns",
    bold_part="Samples")
for bullet in [
    "Enter numbers only — e.g. type  -23  not  -23dB",
    "Negative numbers = amplitude below target = likely PASS",
    "Positive numbers = amplitude above target = review carefully",
]:
    p = doc.add_paragraph(style="List Bullet 2")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(1.5)
    r = p.add_run(bullet); r.font.size = Pt(10); r.font.color.rgb = GREY_TEXT
add_numbered_step(doc, 3, "Save and close the Excel file")
add_numbered_step(doc, 4, "Double-click  Step7_Update_Results.bat",
                  bold_part="Step7_Update_Results.bat")
add_numbered_step(doc, 5, "A window will appear showing progress — wait for it to finish")
add_numbered_step(doc, 6, "Open each report and verify the Pass/Fail results")
add_numbered_step(doc, 7, "Adjust manually if needed")

doc.add_paragraph()

# ── Pass/Fail table ────────────────────────────────────────────────────
add_heading(doc, "Pass / Fail Reference", level=2)
add_pf_table(doc)

doc.add_paragraph()

# ── Something not working ─────────────────────────────────────────────
add_heading(doc, "Something Not Working?", level=2)
add_body(doc,
    "Open a new conversation with Claude Code and describe the problem. "
    "The technical background is in CONTEXT.md in this project folder.")

# ── Save ──────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Created: {OUTPUT}")
